"""Extraction degrades, never crashes."""
from wsbench.extract import extract_text


def test_plain_text(tmp_path):
    p = tmp_path / "out.md"
    p.write_text("# heading\nbody")
    text, note = extract_text(str(p))
    assert text == "# heading\nbody"
    assert note is None


def test_truncation_notes(tmp_path):
    p = tmp_path / "big.txt"
    p.write_text("x" * 100)
    text, note = extract_text(str(p), max_chars=10)
    assert text == "x" * 10
    assert "truncated" in note


def test_unknown_binary_reports_not_extractable(tmp_path):
    p = tmp_path / "blob.zip"
    p.write_bytes(b"\x50\x4b\x03\x04\xff\xfe\x00\x01")
    text, note = extract_text(str(p))
    assert text == ""
    assert "not extractable" in note


def test_docx_roundtrip(tmp_path):
    import docx

    doc = docx.Document()
    doc.add_paragraph("hello docx")
    table = doc.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "a"
    table.rows[0].cells[1].text = "b"
    p = tmp_path / "d.docx"
    doc.save(str(p))

    text, note = extract_text(str(p))
    assert "hello docx" in text
    assert "a | b" in text
    assert note is None


def test_xlsx_roundtrip(tmp_path):
    import openpyxl

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "Data"
    ws.append(["name", "qty"])
    ws.append(["bolt", 42])
    p = tmp_path / "s.xlsx"
    wb.save(str(p))

    text, note = extract_text(str(p))
    assert "# sheet: Data" in text
    assert "bolt\t42" in text
    assert note is None


def test_xlsx_reports_chart_metadata(tmp_path):
    import openpyxl
    from openpyxl.chart import BarChart, Reference

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.append(["item", "amount"])
    ws.append(["a", 3])
    ws.append(["b", 5])
    chart = BarChart()
    chart.type = "bar"  # horizontal
    chart.title = "Spending Comparison"
    chart.add_data(Reference(ws, min_col=2, min_row=1, max_row=3), titles_from_data=True)
    ws.add_chart(chart, "A6")
    p = tmp_path / "c.xlsx"
    wb.save(str(p))

    text, note = extract_text(str(p))
    assert "BarChart (horizontal)" in text
    assert "Spending Comparison" in text
    assert "anchored at" in text
    assert note is None


def test_corrupt_office_file_degrades(tmp_path):
    p = tmp_path / "broken.docx"
    p.write_bytes(b"this is not a zip archive")
    text, note = extract_text(str(p))
    assert text == ""
    assert "extraction failed" in note
